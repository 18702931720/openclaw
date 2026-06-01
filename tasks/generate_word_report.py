#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成 OpenClaw 框架商业化产品竞品分析 Word 文档"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_heading(doc, text, level=1):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(doc, text, bold=False, indent=False):
    """添加段落"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p

def create_table(doc, headers, rows, col_widths=None):
    """创建表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(hdr_cells[i], '2E74B5')
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(10)

    # 数据行
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            row_cells[col_idx].text = str(cell_text)
            row_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in row_cells[col_idx].paragraphs[0].runs:
                run.font.size = Pt(9)
        # 隔行着色
        if row_idx % 2 == 1:
            for cell in row_cells:
                set_cell_shading(cell, 'F2F2F2')

    return table

def main():
    doc = Document()

    # ============ 文档标题 ============
    title = doc.add_heading('OpenClaw 框架商业化产品竞品分析报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 报告信息
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info_para.add_run('报告日期：2026-05-20    |    分析范围：12 款主流商业化产品')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()

    # ============ 一、市场概览 ============
    add_heading(doc, '一、市场概览', 1)
    p = doc.add_paragraph()
    run = p.add_run('OpenClaw 是一款开源 AI Agent 框架，核心能力是模拟人类交互行为完成 Web/APP 界面的自动化操作——发邮件、写代码、整理文件、远程控制电脑等。2026 年初以来，MiniMax、月之暗面、智谱、腾讯、火山引擎、网易有道、阿里云等大厂纷纷基于 OpenClaw 推出简化版或云托管版本，将门槛从极客专属降低到普通用户可操作的范围。')
    run.font.size = Pt(11)

    p2 = doc.add_paragraph()
    run2 = p2.add_run('市场格局：形成三大阵营——互联网大厂系（腾讯/阿里/字节/百度/网易）、AI 厂商系（MiniMax/月之暗面/智谱）、垂直专业版（QeeClaw/RetailClaw）。')
    run2.font.size = Pt(11)

    # ============ 二、产品矩阵 ============
    add_heading(doc, '二、产品矩阵与定位对比', 1)

    # 2.1 互联网大厂系
    add_heading(doc, '2.1 互联网大厂系（6 款）', 2)
    headers1 = ['产品', '厂商', '定位人群', '核心亮点', '定价']
    rows1 = [
        ['QClaw', '腾讯', '重度微信用户', '微信/QQ 直连，远程控制电脑', '内测免费/低订阅'],
        ['WorkBuddy', '腾讯', '企业团队', '企业级 AI 助手，审计日志+权限控制', '企业订阅'],
        ['CoPaw', '阿里云', '电商/大企业', '钉钉/飞书/QQ 全家桶集成', '商业订阅'],
        ['钉钉悟空', '钉钉', '各类企业', '智能办公助手，会议管理+文档协作', '实惠'],
        ['DuClaw', '百度', '信息检索需求者', '搜索+知识图谱', '按量计费'],
        ['ArkClaw', '火山引擎（字节）', '飞书团队/企业', '飞书深度集成，企业级高并发', 'Lite 首月 9.9 元起'],
    ]
    create_table(doc, headers1, rows1)
    doc.add_paragraph()

    # 2.2 AI 厂商系
    add_heading(doc, '2.2 AI 厂商系（3 款）', 2)
    headers2 = ['产品', '厂商', '定位人群', '核心亮点', '定价']
    rows2 = [
        ['MaxClaw', 'MiniMax', '预算党/小白', '性价比之选，云端一键，MoE 模型', '~39 元/月'],
        ['KimiClaw', '月之暗面', '知识工作者', '40GB 云存储，超长上下文处理', '199 元/月'],
        ['AutoClaw', '智谱 AI', '隐私本地党', '本地一键部署，浏览器自动化强', '免费+积分'],
    ]
    create_table(doc, headers2, rows2)
    doc.add_paragraph()

    # 2.3 垂直专业版
    add_heading(doc, '2.3 垂直专业版（2 款）', 2)
    headers3 = ['产品', '厂商', '定位人群', '核心亮点']
    rows3 = [
        ['QeeClaw', '第三方', '企业客户', '企业超级秘书，私有化部署，一体机'],
        ['RetailClaw', '第三方', '亚马逊品牌', 'Vendor Central 运营自动化'],
    ]
    create_table(doc, headers3, rows3)
    doc.add_paragraph()

    # ============ 三、核心维度对比 ============
    add_heading(doc, '三、核心维度对比分析', 1)

    add_heading(doc, '3.1 部署方式', 2)
    headers_d = ['部署模式', '代表产品']
    rows_d = [
        ['云端一键', 'MaxClaw、KimiClaw、ArkClaw'],
        ['本地客户端', 'QClaw、AutoClaw、LobsterAI'],
        ['混合部署', 'CoPaw、WorkBuddy'],
        ['开源自托管', 'OpenClaw 原版'],
    ]
    create_table(doc, headers_d, rows_d)
    doc.add_paragraph()

    add_heading(doc, '3.2 定价区间', 2)
    headers_p = ['价格层级', '产品']
    rows_p = [
        ['免费', 'OpenClaw 原版、AutoClaw（免费额度）'],
        ['低价（<50 元/月）', 'MaxClaw（~39 元/月）、ArkClaw Lite（9.9 元首月）'],
        ['中价（50-200 元/月）', 'KimiClaw（199 元）、LobsterAI'],
        ['企业定价（面议）', 'QClaw、WorkBuddy、CoPaw、QeeClaw、RetailClaw'],
    ]
    create_table(doc, headers_p, rows_p)
    doc.add_paragraph()

    add_heading(doc, '3.3 生态集成能力', 2)
    headers_e = ['集成深度', '产品']
    rows_e = [
        ['微信/QQ 生态', 'QClaw（最强）、CoPaw'],
        ['飞书生态', 'ArkClaw（深度集成）、WorkBuddy'],
        ['钉钉生态', '钉钉悟空、CoPaw'],
        ['阿里云生态', 'CoPaw（全家桶整合）'],
        ['独立平台', 'MaxClaw、KimiClaw、AutoClaw（跨平台）'],
    ]
    create_table(doc, headers_e, rows_e)
    doc.add_paragraph()

    # ============ 四、逐款深度分析 ============
    add_heading(doc, '四、逐款产品深度分析', 1)

    products = [
        {
            'name': '4.1 QClaw（腾讯）',
            'subtitle': '微信通道最强玩家',
            'position': '"你的 AI 员工，不是 AI 玩具"',
            'abilities': [
                '接入微信/QQ，直接在微信里操控',
                '能管理邮件、整理文件、发消息、写代码、搜索网络',
                '替代人工点击和录入，真正的任务执行',
                '支持远程控制电脑',
            ],
            'pros': ['腾讯生态加持，微信通道无可替代', 'OpenClaw 框架成熟，技能生态完善（5000+ 技能）', 'C 端体验最好，普通用户零门槛'],
            'cons': ['目前主要面向个人/中小企业', '企业级安全合规能力待验证', 'Mac 版本仍在内测'],
            'summary': '微信重度用户首选，SaaS 版对个人用户友好，企业版仍在打磨。',
        },
        {
            'name': '4.2 WorkBuddy（腾讯）',
            'subtitle': '企业级安全版',
            'position': '企业办公 AI 助手，强调安全与管控',
            'abilities': [
                '企业 IM 深度集成',
                '审计日志+权限控制',
                '多租户隔离架构',
                'SSO、RBAC 企业级安全',
            ],
            'pros': ['腾讯安全体系背书', '适合对数据安全要求极高的大厂', '企业管理后台完善'],
            'cons': ['定价面向企业，具体价格不透明', '功能相对保守，创新速度可能慢于创业公司'],
            'summary': '适合对数据安全有刚性要求的大企业，但创新速度不如创业公司。',
        },
        {
            'name': '4.3 ArkClaw（火山引擎/字节）',
            'subtitle': '飞书生态企业级方案',
            'position': '飞书团队企业级 AI Agent',
            'abilities': [
                '与飞书无缝集成',
                '高并发稳定运行（99.99% 可用性）',
                'VPC 隔离、等保合规',
                '支持大规模团队部署',
            ],
            'pros': ['字节跳动技术背书', '飞书生态天然优势', '企业级稳定性'],
            'cons': ['过度依赖飞书生态，非飞书用户门槛高', '价格相对较高'],
            'summary': '飞书团队首选，非飞书用户不建议。',
        },
        {
            'name': '4.4 CoPaw（阿里云）',
            'subtitle': '全家桶整合方案',
            'position': '阿里云生态整合，企业上云首选',
            'abilities': [
                '钉钉/飞书/QQ 全家桶原生集成',
                '电商运营全流程支持',
                '商业订阅模式',
            ],
            'pros': ['阿里系软件栈无缝衔接', '电商场景成熟', '企业级能力完善'],
            'cons': ['生态绑定严重，非阿里系用户不友好', '价格不透明，决策成本高'],
            'summary': '已有阿里系软件栈的企业首选，其他企业慎入。',
        },
        {
            'name': '4.5 MaxClaw（MiniMax）',
            'subtitle': '性价比之选',
            'position': '普通用户的零门槛 AI Agent',
            'abilities': [
                '云端一键部署，5 分钟上手',
                '低成本 MoE 模型',
                '专家预置，开箱即用',
            ],
            'pros': ['价格最低（~39 元/月）', '上手难度最低', 'MiniMax 技术背书'],
            'cons': ['功能相对基础，不适合复杂场景', '定制化能力弱'],
            'summary': '预算有限的小白用户首选，功能深度不足。',
        },
        {
            'name': '4.6 KimiClaw（月之暗面）',
            'subtitle': '知识工作者利器',
            'position': '超长上下文处理，专注知识管理',
            'abilities': [
                '40GB 云存储',
                '超长上下文（支持长文档处理）',
                'Kimi 大模型驱动',
            ],
            'pros': ['长文档处理能力最强', '知识管理场景优化', '上下文记忆深度好'],
            'cons': ['价格较高（199 元/月）', '办公自动化能力弱于 QClaw'],
            'summary': '重度知识工作者首选，办公自动化场景不如 QClaw。',
        },
        {
            'name': '4.7 AutoClaw（智谱 AI）',
            'subtitle': '隐私本地党首选',
            'position': '本地部署，隐私安全',
            'abilities': [
                '一键本地部署',
                '浏览器自动化能力强',
                '免费额度+积分模式',
                '50+ 预置技能',
            ],
            'pros': ['数据不上云，隐私最高', '免费额度友好', '智谱 AI 技术背书'],
            'cons': ['本地部署有运维成本', '云端协同能力弱'],
            'summary': '隐私敏感用户首选，但运维成本较高。',
        },
        {
            'name': '4.8 钉钉悟空（钉钉）',
            'subtitle': '智能办公入门选择',
            'position': '钉钉用户的智能办公助手',
            'abilities': [
                '会议管理+文档协作+日程安排',
                '钉钉深度集成',
                '价格实惠',
            ],
            'pros': ['钉钉生态天然优势', '价格亲民', '办公基础功能完善'],
            'cons': ['AI 能力深度不如专业 Agent 产品', '定位偏辅助工具而非数字员工'],
            'summary': '钉钉用户入门首选，但 AI 能力天花板较低。',
        },
    ]

    for p in products:
        # 产品名
        h = doc.add_heading(p['name'], 2)
        # 副标题
        p_sub = doc.add_paragraph()
        run_b = p_sub.add_run(f"定位：{p['subtitle']}（{p['position']}）")
        run_b.bold = True
        run_b.font.size = Pt(11)

        # 关键能力
        doc.add_paragraph('关键能力：', style='List Bullet').runs[0].bold = True
        for a in p['abilities']:
            doc.add_paragraph(a, style='List Bullet')

        # 优势
        doc.add_paragraph().add_run('优势：').bold = True
        for pro in p['pros']:
            doc.add_paragraph(pro, style='List Bullet')

        # 劣势
        doc.add_paragraph().add_run('劣势：').bold = True
        for con in p['cons']:
            doc.add_paragraph(con, style='List Bullet')

        # 一句话总结
        p_sum = doc.add_paragraph()
        run_sum = p_sum.add_run(f"一句话总结：{p['summary']}")
        run_sum.italic = True
        run_sum.font.size = Pt(10)
        run_sum.font.color.rgb = RGBColor(70, 130, 180)

        doc.add_paragraph()

    # ============ 五、竞争格局分析 ============
    add_heading(doc, '五、竞争格局分析', 1)

    add_heading(doc, '5.1 赛道定位矩阵', 2)
    headers_m = ['维度', 'C端/小白用户', 'B端/企业用户']
    rows_m = [
        ['泛化平台', 'MaxClaw、KimiClaw', 'WorkBuddy、CoPaw'],
        ['垂直场景', 'QClaw（微信党）', 'ArkClaw（飞书）、QeeClaw（企业）、RetailClaw（电商）'],
        ['技术底座', 'OpenClaw 原版', 'AutoClaw（本地）、LobsterAI（办公）'],
    ]
    create_table(doc, headers_m, rows_m)
    doc.add_paragraph()

    add_heading(doc, '5.2 核心竞争要素', 2)
    factors = [
        ('通道优势', 'QClaw 凭微信通道建立差异化壁垒，其他产品难以复制'),
        ('生态绑定', 'ArkClaw（飞书）、CoPaw（阿里全家桶）依赖生态，限制跨平台能力'),
        ('价格分层', 'MaxClaw（39 元）vs KimiClaw（199 元）形成鲜明价格带区分'),
        ('隐私与便捷', 'AutoClaw（隐私优先）vs QClaw（便捷优先），面向不同需求用户'),
    ]
    for title, desc in factors:
        p = doc.add_paragraph()
        run_t = p.add_run(f"{title}：")
        run_t.bold = True
        p.add_run(desc)
    doc.add_paragraph()

    # ============ 六、关键洞察与建议 ============
    add_heading(doc, '六、关键洞察与建议', 1)

    add_heading(doc, '6.1 市场机会', 2)
    opportunities = [
        ('AI Agent 大爆发', '2026 年是 AI Agent 元年，OpenClaw 衍生产品进入爆发期'),
        ('C 端蓝海', '小白用户对"零门槛 AI 助手"需求旺盛，但产品供给不足'),
        ('企业级需求', '安全合规+私有化部署成为大企业刚需，但价格敏感'),
    ]
    for t, d in opportunities:
        p = doc.add_paragraph()
        run_t = p.add_run(f"• {t}：")
        run_t.bold = True
        p.add_run(d)

    doc.add_paragraph()

    add_heading(doc, '6.2 竞争威胁', 2)
    threats = [
        ('QClaw 最具威胁', '腾讯品牌+微信通道+C 端体验，对个人用户吸引力最强'),
        ('大厂生态绑定', '飞书/钉钉/阿里云用户被生态产品锁定，跨平台获客难'),
        ('同质化风险', '大部分产品功能相似，差异化空间有限'),
    ]
    for t, d in threats:
        p = doc.add_paragraph()
        run_t = p.add_run(f"• {t}：")
        run_t.bold = True
        p.add_run(d)

    doc.add_paragraph()

    add_heading(doc, '6.3 潜在机会点', 2)
    headers_o = ['机会点', '说明']
    rows_o = [
        ['企业微信通道', 'QClaw 占据微信，企微可能是空白或弱竞争区'],
        ['垂直行业 Agent', '金融/医疗/政务等行业定制化需求未被充分满足'],
        ['私有化+轻量化', '企业级安全+小白易用性的结合点'],
        ['多通道整合', '微信+飞书+钉钉多平台统一管理'],
    ]
    create_table(doc, headers_o, rows_o)
    doc.add_paragraph()

    # ============ 七、总结 ============
    add_heading(doc, '七、总结', 1)
    p = doc.add_paragraph()
    run = p.add_run('OpenClaw 框架衍生产品已形成完整的市场矩阵，从免费开源到企业订阅，从个人小白到大型企业均有覆盖。')
    run.font.size = Pt(11)

    conclusions = [
        'QClaw 是当前最成熟的 C 端产品，凭借微信通道和腾讯品牌建立独特优势',
        'ArkClaw 和 CoPaw 依赖大厂生态，在各自细分场景有竞争力',
        'MaxClaw 和 KimiClaw 主打低价和知识管理，填补小白市场空白',
        '企业微信通道可能是差异化机会，垂直行业定制和轻量化企业版也是值得探索的方向',
    ]
    for c in conclusions:
        doc.add_paragraph(f"• {c}")

    doc.add_paragraph()

    # ============ 附录：选型决策树 ============
    add_heading(doc, '附录：产品选型决策树', 1)
    decisions = [
        ('隐私第一 →', 'AutoClaw / LobsterAI（本地部署）'),
        ('微信重度 →', 'QClaw（最强）'),
        ('飞书团队 →', 'ArkClaw'),
        ('预算紧张 →', 'MaxClaw'),
        ('知识工作 →', 'KimiClaw'),
        ('大企业/高安全 →', 'WorkBuddy / CoPaw'),
        ('电商垂直 →', 'RetailClaw'),
        ('开源极客 →', 'OpenClaw 原版'),
    ]
    for d in decisions:
        p = doc.add_paragraph()
        run1 = p.add_run(d[0])
        run1.bold = True
        p.add_run(d[1])

    # 文档备注
    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_note = note.add_run('_本报告基于 2026 年 5 月公开信息整理，定价和功能可能有变动，仅供参考。_')
    run_note.font.size = Pt(9)
    run_note.font.color.rgb = RGBColor(128, 128, 128)

    # 保存
    output_path = '/home/admin/.openclaw/workspace/tasks/OpenClaw竞品分析报告_20260520.docx'
    doc.save(output_path)
    print(f'文档已生成：{output_path}')
    return output_path

if __name__ == '__main__':
    main()